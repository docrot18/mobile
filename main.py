# Импорт всех классов
from docx.shared import Pt
from kivy.app import App
from kivy.lang import Builder
from kivy.uix.anchorlayout import AnchorLayout
from kivy.uix.gridlayout import GridLayout
from kivy.uix.label import Label
from kivy.uix.button import Button
from kivy.uix.scrollview import ScrollView
from kivy.uix.textinput import TextInput
from kivy.uix.screenmanager import ScreenManager, Screen
from kivy.uix.boxlayout import BoxLayout
import openpyxl
import docx
import sqlite3
from docx2pdf import convert

from kivy.core.window import Window

# Глобальные настройки
Window.clearcolor = (180 / 255, 180 / 255, 180 / 255, 1)
Window.title = "ОКЦ"


class Auth(Screen):
    Builder.load_file("auth.kv")

    def login(self, login, password):

        conn = sqlite3.connect('auth.db')
        c = conn.cursor()
        test = c.execute(f""" SELECT * FROM users WHERE login = '{login}' AND password = '{password}'""").fetchone()

        if test is None:
            self.ids['error'].text = 'Логин или пароль введены неверно'
        else:

            return self.secondScreen()
            conn.commit()
        conn.close()

    def secondScreen(self, *args):
        self.manager.current = 'Second'


class MainWindow(Screen):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self.name = 'Second'
        conn = sqlite3.connect('auth.db')
        c = conn.cursor()
        c.execute(""" SELECT * FROM products""")

        mainbox = BoxLayout(orientation="vertical")
        containerMain = BoxLayout(orientation="vertical", size_hint_y=None)
        scrollbar = ScrollView(do_scroll_y=True, do_scroll_x=False)
        global nameDict, quantDict, quantDictTemp
        nameDict = dict()
        quantDictTemp = dict()
        quantDict = dict()

        global minusDict, plusDict
        minusDict = dict()
        plusDict = dict()
        infoTabs = BoxLayout(size_hint=(1, 0.15))
        submit = Button(text="Сохранить заказ", size_hint=(1, 0.15), on_press=self.fouthScreen)
        adminTab = Button(text="Редактировать позиции", on_press=self.thirdScreen)
        orderTab = Button(text="Создать заказ", disabled=True)

        for id, nameDB, quant in c.fetchall():
            container = GridLayout(cols=4, spacing=5)
            nameDict[id] = Label(width=mainbox.width / 100.0 * 70.0)

            nameDict[id].text = nameDB
            container2 = GridLayout(cols=4, spacing=5, size_hint=(1, 1))
            quantDictTemp[id] = TextInput(width=(mainbox.width / 10.0))
            quantityBox = BoxLayout(orientation="vertical")
            quantityBoxMax = BoxLayout()
            quantityMaxText = Label(text="Всего: ",width=(mainbox.width / 10.0))
            quantDict[id] = Label(text=str(quant),width=(mainbox.width / 10.0))
            quantDictTemp[id].text = "0"
            quantDictTemp[id].id = quant
            minusDict[id] = Button(text="-", font_size=20, on_press=lambda minus: self.minusQuantity(minus))
            plusDict[id] = Button(text="+",font_size=20, on_press=lambda plus: self.plusQuantity(plus))
            container2.add_widget(minusDict[id])
            quantityBox.add_widget(quantDictTemp[id])
            quantityBoxMax.add_widget(quantityMaxText)
            quantityBoxMax.add_widget(quantDict[id])
            quantityBox.add_widget(quantityBoxMax)

            container2.add_widget(quantityBox)

            container2.add_widget(plusDict[id])
            container.add_widget(nameDict[id])
            container.add_widget(container2)
            containerMain.add_widget(container)

        infoTabs.add_widget(orderTab)
        infoTabs.add_widget(adminTab)
        mainbox.add_widget(infoTabs)
        scrollbar.add_widget(containerMain)
        height = len(containerMain.children)
        containerMain.height = height * 60
        mainbox.add_widget(scrollbar)
        mainbox.add_widget(submit)
        self.add_widget(mainbox)

    def minusQuantity(self, obj):
        global labels
        if obj.parent.children[1].children[1].text == "0":
            return 0
        else:
            temp = obj.parent.children[1].children[1].text
            result = int(temp) - 1
            obj.parent.children[1].children[1].text = str(result)

    def plusQuantity(self, obj):
        global labels
        if obj.parent.children[1].children[1].text == obj.parent.children[1].children[0].children[0].text:
            return 0
        else:
            temp = obj.parent.children[1].children[1].text
            result = int(temp) + 1
            obj.parent.children[1].children[1].text = str(result)

    def thirdScreen(self, *args):
        self.manager.current = 'Third'


    def fouthScreen(self, *args):
        self.manager.current = 'Fouth'




class AdminWindow(Screen):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self.name = 'Third'
        conn = sqlite3.connect('auth.db')
        c = conn.cursor()
        c.execute(""" SELECT * FROM products""")
        mainbox = BoxLayout(orientation="vertical")
        containerMain = BoxLayout(orientation="vertical", size_hint_y=None)
        global minusDictAdmin, plusDictAdmin, nameDictAdmin, quantDictAdmin, idDictAdmin
        minusDictAdmin = dict()
        nameDictAdmin = dict()
        idDictAdmin = dict()
        quantDictAdmin = dict()
        plusDictAdmin = dict()
        info = BoxLayout(size_hint=(1, 0.15))
        adminTab = Button(text="Редактировать позиции", disabled=True)
        orderTab = Button(text="Создать заказ", on_press=self.secondScreen)
        BottomBox = BoxLayout(size_hint=(1, 0.15))
        submit = Button(text="Сохранить", size_hint=(1, 1), on_press=self.saveDB)
        add = Button(text="Добавить", size_hint=(1, 1), on_press=self.addRow)

        scrollbar = ScrollView(do_scroll_y=True, do_scroll_x=False)
        for id, nameDB, quant in c.fetchall():
            container = GridLayout(cols=4, spacing=5)
            nameDictAdmin[id] = TextInput(width=mainbox.width / 100.0 * 70.0)
            nameDictAdmin[id].text = nameDB
            container2 = GridLayout(cols=3, spacing=5)
            quantDictAdmin[id] = TextInput(width=(mainbox.width / 10.0))
            quantDictAdmin[id].text = str(quant)
            quantDictAdmin[id].id = str(id)
            idDictAdmin[id] = id
            minusDictAdmin[id] = Button(text="-", font_size=20, width=mainbox.width / 10.0,
                                   on_press=lambda minus: self.minusQuantity(minus))
            plusDictAdmin[id] = Button(text="+", font_size=20, width=mainbox.width / 10.0, on_press=lambda plus: self.plusQuantity(plus))
            container2.add_widget(minusDictAdmin[id])
            container2.add_widget(quantDictAdmin[id])
            container2.add_widget(plusDictAdmin[id])
            container.add_widget(nameDictAdmin[id])
            container.add_widget(container2)
            containerMain.add_widget(container)

        info.add_widget(orderTab)
        info.add_widget(adminTab)
        mainbox.add_widget(info)
        scrollbar.add_widget(containerMain)
        height = len(containerMain.children)
        containerMain.height = height * 60
        mainbox.add_widget(scrollbar)
        BottomBox.add_widget(submit)
        BottomBox.add_widget(add)
        mainbox.add_widget(BottomBox)

        self.add_widget(mainbox)

    def minusQuantity(self, obj):
        global labels
        if obj.parent.children[1].text == "0":
            return 0
        else:
            temp = obj.parent.children[1].text
            result = int(temp) - 1
            obj.parent.children[1].text = str(result)

    def plusQuantity(self, obj):
        global labels
        temp = obj.parent.children[1].text
        result = int(temp) + 1
        obj.parent.children[1].text = str(result)

    def secondScreen(self, *args):
        self.manager.current = 'Second'

    def saveDB(self, *args):
        conn = sqlite3.connect('auth.db')
        c = conn.cursor()
        c.execute(""" SELECT * FROM products""")
        count = 0
        for id, nameDB, quant in c.fetchall():
            if nameDB != nameDictAdmin[id].text:
                c.execute(f""" UPDATE products SET name = REPLACE('{nameDB}', '{nameDB}', '{nameDictAdmin[id].text}') WHERE id = '{id}'""")
                nameDict[id].text = nameDictAdmin[id].text
            if quant != quantDictAdmin[id].text:
                c.execute(f""" UPDATE products SET quantity = REPLACE('{quant}', '{quant}', '{quantDictAdmin[id].text}') WHERE id = '{id}'""")
                quantDict[id].text = quantDictAdmin[id].text

            conn.commit()
            count+=1

    def addRow(self, *args):
        conn = sqlite3.connect('auth.db')
        c = conn.cursor()
        c.execute(f""" INSERT INTO products (name, quantity) VALUES ("Новая Позиция", "1")""")
        conn.commit()
        conn.close()


class SaveWindow(Screen):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self.name = 'Fouth'
        conn = sqlite3.connect('auth.db')
        c = conn.cursor()
        c.execute(""" SELECT * FROM products""")
        mainbox = AnchorLayout()
        container = BoxLayout(orientation="vertical", size_hint=(0.5, 0.5))
        excel = Button(text="MS Excel", on_press=self.Excel)
        word = Button(text="MS Word",on_press=self.Word)
        pdf = Button(text="PDF", on_press=self.PDF)
        container.add_widget(excel)
        container.add_widget(word)
        container.add_widget(pdf)
        mainbox.add_widget(container)
        back = AnchorLayout(anchor_x="left", anchor_y="top")
        backBtn = Button(text="Назад", size_hint=(0.2, 0.14), on_press=self.secondScreen)
        back.add_widget(backBtn)
        mainbox.add_widget(back)
        self.add_widget(mainbox)

    def secondScreen(self, *args):
        self.manager.current = 'Second'

    def Excel(self, *args):
        conn = sqlite3.connect('auth.db')
        c = conn.cursor()
        c.execute(""" SELECT * FROM products""")
        global nameDict, quantDict, quantDictTemp
        string = ""
        count = 1
        wb = openpyxl.Workbook()
        work_sheet = wb.create_sheet('Заказ', 0)
        for id, nameDB, quant in c.fetchall():
            if quantDictTemp[id].text != "0":
                work_sheet.append([nameDict[id].text, quantDictTemp[id].text])
        wb.save("Заказ.xlsx")
        wb.close()
        c.execute(f""" INSERT INTO orders (body) VALUES ('{string}')""")
        conn.commit()
        conn.close()




    def Word(self, *args):
        conn = sqlite3.connect('auth.db')
        c = conn.cursor()
        c.execute(""" SELECT * FROM products""")
        global nameDict, quantDict, quantDictTemp
        string=""
        count=1
        for id, nameDB, quant in c.fetchall():
            if quantDictTemp[id].text != "0":
                string+=str(count)+". "+nameDict[id].text+ ". Количество: "+quantDictTemp[id].text + "\n"
                count+=1
        c.execute(f""" INSERT INTO orders (body) VALUES ('{string}')""")
        conn.commit()
        conn.close()
        doc = docx.Document()
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(16)
        doc.add_paragraph(string)
        doc.save("Заказ.docx")

    def PDF(self, *args):
        conn = sqlite3.connect('auth.db')
        c = conn.cursor()
        c.execute(""" SELECT * FROM products""")
        global nameDict, quantDict, quantDictTemp
        string = ""
        count = 1
        for id, nameDB, quant in c.fetchall():
            if quantDictTemp[id].text != "0":
                string += str(count) + ". " + nameDict[id].text + ". Количество: " + quantDictTemp[id].text + "\n"
                count += 1
        c.execute(f""" INSERT INTO orders (body) VALUES ('{string}')""")
        conn.commit()
        conn.close()
        doc = docx.Document()
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(16)
        doc.add_paragraph(string)
        doc.save("Заказ.docx")
        convert("Заказ.docx")


class MyApp(App):

    # Основной метод для построения программы
    def build(self):
        conn = sqlite3.connect('auth.db')
        c = conn.cursor()
        c.execute("""CREATE TABLE if not exists users(
                    id INTEGER,
        			login TEXT,
        			password TEXT)
        		 """)

        c.execute("""CREATE TABLE if not exists products(
                    id INTEGER,
        			name TEXT,
        			quantity INTEGER)
        		 """)

        c.execute("""CREATE TABLE if not exists orders(
                    id INTEGER,
        			body TEXT)
        		 """)

        conn.commit()
        conn.close()
        sm.add_widget(Auth())
        sm.add_widget(MainWindow())
        sm.add_widget(AdminWindow())
        sm.add_widget(SaveWindow())
        return sm


sm = ScreenManager()

# Запуск проекта
if __name__ == "__main__":
    MyApp().run()
